import os
import re
from docx import Document
from datetime import time
from slugify import slugify
from schedule.models import Group, Teacher, Classroom, Subject, Lesson, TimeSlot, GroupDenomination


DAYS_ORDER = ['ПОНЕДЕЛЬНИК', 'ВТОРНИК', 'СРЕДА', 'ЧЕТВЕРГ', 'ПЯТНИЦА', 'СУББОТА']
DAYS_MAP = {'понедельник': 1, 'вторник': 2, 'среда': 3, 'четверг': 4, 'пятница': 5, 'суббота': 6}


def extract_course_from_doc(doc: Document) -> int:
    """Извлекает курс из документа"""
    for para in doc.paragraphs:
        text = para.text.strip().lower()
        match = re.search(r'(\d+)\s*курс', text)
        if match:
            return int(match.group(1))
    return 1


def find_schedule_table(doc: Document):
    """Находит таблицу с расписанием по характерным признакам"""
    for table in doc.tables:
        if len(table.rows) < 3:
            continue
        first_row = [c.text.strip().lower() for c in table.rows[0].cells]
        if 'время' in first_row and 'дисциплины' in first_row:
            return table
    
    for table in doc.tables:
        for row in table.rows[:5]:
            row_text = ' '.join([cell.text for cell in row.cells]).lower()
            if 'понедельник' in row_text and 'вторник' in row_text:
                return table
    
    return doc.tables[0] if doc.tables else None


def parse_docx(file_path: str, clear: bool = False, group=None):
    """Импорт расписания из docx файла по указанному пути"""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f'Файл не найден: {file_path}')

    if clear and group:
        Lesson.objects.filter(group=group).delete()

    doc = Document(file_path)
    return _parse_doc(doc, group=group)


def parse_docx_file(file_obj, clear: bool = False, group=None):
    """Импорт расписания из файлового объекта (для загрузки через веб)"""
    if clear and group:
        Lesson.objects.filter(group=group).delete()

    doc = Document(file_obj)
    return _parse_doc(doc, group=group)





def _parse_doc(doc: Document, group=None) -> dict:
    """Парсит документ docx"""
    schedule_table = find_schedule_table(doc)
    
    if not schedule_table:
        return {'lessons': 0, 'errors': ['Таблица с расписанием не найдена']}

    if not group:
        return {'lessons': 0, 'errors': ['Группа не выбрана']}
    
    lessons_created = 0
    errors = []
    current_day = None
    
    for row_idx, row in enumerate(schedule_table.rows):
        cells = row.cells
        if not cells:
            continue

        all_cells_text = [c.text.strip() for c in cells]
        all_cells_lower = [t.lower() for t in all_cells_text]

        day_found = False
        for i, cell_text in enumerate(all_cells_lower):
            for day in DAYS_ORDER:
                if day in cell_text.upper() and len(cell_text) < 20:
                    current_day = DAYS_MAP[day.lower()]
                    day_found = True
                    break
            if day_found:
                break

        if day_found:
            continue

        first_cell = all_cells_text[0]
        time_pattern = r'\d{1,2}\.\d{2}-\d{1,2}\.\d{2}'
        time_match = re.search(time_pattern, first_cell)
        
        if not time_match:
            continue

        time_str = first_cell.strip()
        
        time_slot = _get_or_create_time_slot(time_str)
        
        discipline = all_cells_text[1] if len(cells) > 1 else ""
        group_info = all_cells_text[2] if len(cells) > 2 else ""
        lesson_type_str = all_cells_text[3] if len(cells) > 3 else ""
        teacher_name = all_cells_text[4] if len(cells) > 4 else ""
        room_name = all_cells_text[5] if len(cells) > 5 else ""

        if not current_day or not discipline:
            continue

        week_type = 'BOTH'
        
        # Извлекаем номер группы из названия (например "1 группа" -> 1)
        group_num = 0
        match = re.search(r'(\d+)', group.name)
        if match:
            group_num = int(match.group(1))
        
        denomination = _parse_denomination(group_info, group.name, group_num)
        
        if denomination is None and group_info.strip():
            # Если есть текст но denomination = None (например число = номер группы) - это "вся группа", добавляем
            pass
        
        lesson_count = _create_lesson(
            group=group,
            discipline=discipline,
            time_slot=time_slot,
            teacher_name=teacher_name,
            room_name=room_name,
            lesson_type_str=lesson_type_str,
            day=current_day,
            week_type=week_type,
            errors=errors,
            time_str=time_str,
            denomination=denomination
        )
        lessons_created += lesson_count

    return {'lessons': lessons_created, 'errors': errors}


def _create_lesson(group, discipline, time_slot, teacher_name, room_name, lesson_type_str, day, week_type, errors, time_str, denomination=None):
    """Создает занятие и возвращает количество созданных (0 или 1)"""
    subject, _ = Subject.objects.get_or_create(name=discipline)

    teacher = None
    if teacher_name:
        clean_teacher = _clean_teacher_name(teacher_name)
        if clean_teacher:
            slug = slugify(clean_teacher)
            defaults = {'slug': slug} if slug else {}
            teacher, created = Teacher.objects.get_or_create(
                name=clean_teacher,
                defaults=defaults
            )
            if not created and teacher.slug is None and slug:
                teacher.slug = slug
                teacher.save()

    classroom = None
    if room_name:
        slug = slugify(room_name)
        if slug:
            classroom, created = Classroom.objects.get_or_create(
                name=room_name,
                defaults={'slug': slug}
            )
            if not created and classroom.slug is None:
                classroom.slug = slug
                classroom.save()

    lesson_type = _parse_lesson_type(lesson_type_str)

    existing = Lesson.objects.filter(
        group=group,
        subject=subject,
        time_slot=time_slot,
        day_of_week=day,
        week_type=week_type,
        denomination=denomination
    ).first()

    if existing:
        return 0

    try:
        Lesson.objects.create(
            group=group,
            subject=subject,
            teacher=teacher,
            classroom=classroom,
            time_slot=time_slot,
            day_of_week=day,
            week_type=week_type,
            lesson_type=lesson_type,
            denomination=denomination
        )
        return 1
    except Exception as e:
        errors.append(f'{time_str} {discipline}: {str(e)}')
        return 0


def _parse_time(time_str: str):
    """Парсит время в формате XX.XX-XX.XX"""
    pattern = r'(\d{2})\.(\d{2})-'
    match = re.search(pattern, time_str)
    
    if match:
        hour = int(match.group(1))
        minute = int(match.group(2))
        return time(hour, minute)
    
    return time(9, 0)


def _get_or_create_time_slot(time_str: str):
    """Получает или создает TimeSlot на основе строки времени"""
    pattern = r'(\d{1,2})\.(\d{2})-(\d{1,2})\.(\d{2})'
    match = re.search(pattern, time_str)
    
    if match:
        start = time(int(match.group(1)), int(match.group(2)))
        end = time(int(match.group(3)), int(match.group(4)))
        
        ts = TimeSlot.objects.filter(start_time=start, end_time=end).first()
        if ts:
            return ts
        
        name = f"{int(match.group(1)):02d}.{match.group(2)}-{int(match.group(3)):02d}.{match.group(4)}"
        return TimeSlot.objects.create(name=name, start_time=start, end_time=end)
    
    return None


def _parse_group_name(group_info: str) -> str:
    """Парсит название группы из строки"""
    if not group_info:
        return None
    
    group_info = group_info.replace('\n', ' ').strip()
    
    if 'все студенты' in group_info.lower():
        return 'ALL'
    
    match = re.match(r'(\d+)\s*группа', group_info.lower())
    if match:
        return f"{match.group(1)} группа"
    
    match = re.match(r'(\d+)\s*подгруппа', group_info.lower())
    if match:
        return f"{match.group(1)} подгруппа"
    
    if group_info:
        return group_info.split()[0] if group_info.split() else None
    
    return None


def _parse_group_from_first_row(group_info: str) -> str:
    """Парсит номер группы из первой строки (информационный час)"""
    if not group_info:
        return None
    
    group_info = group_info.replace('\n', ' ').strip()
    
    match = re.match(r'(\d+)', group_info)
    if match:
        return f"{match.group(1)} группа"
    
    if group_info:
        return group_info.split()[0] if group_info.split() else None
    
    return None


def _parse_denomination(group_info: str, group_name: str, group_num: int) -> GroupDenomination | None:
    """Парсит вид группы из третьей колонки
    
    Правила:
    - Число = номер группы -> вся группа (None)
    - Число != номер группы -> смотрим слово после числа:
      - "подгруппа" -> подгруппа
      - "группа" -> группа  
      - другое -> подгруппа
    """
    if not group_info:
        return None
    
    info = group_info.strip().lower()
    if not info:
        return None
    
    # Ищем число и слово после него
    match = re.search(r'(\d+)\s*(\w+)', info)
    if not match:
        return None
    
    number = int(match.group(1))
    word_after_number = match.group(2)
    
    # Число = номер группы -> вся группа
    if number == group_num:
        return None
    
    # Определяем тип по слову после числа
    if 'подгруппа' in word_after_number or 'п' in word_after_number[:1]:
        # Подгруппа
        denom, _ = GroupDenomination.objects.get_or_create(
            name=f"Подгруппа {number}",
            type='SUBGROUP',
            defaults={'name': f"Подгруппа {number}", 'type': 'SUBGROUP'}
        )
        return denom
    elif 'группа' in word_after_number or 'г' in word_after_number[:1]:
        # Группа
        denom, _ = GroupDenomination.objects.get_or_create(
            name=f"Группа {number}",
            type='GROUP',
            defaults={'name': f"Группа {number}", 'type': 'GROUP'}
        )
        return denom
    else:
        # По умолчанию подгруппа
        denom, _ = GroupDenomination.objects.get_or_create(
            name=f"Подгруппа {number}",
            type='SUBGROUP',
            defaults={'name': f"Подгруппа {number}", 'type': 'SUBGROUP'}
        )
        return denom


def _clean_teacher_name(teacher_name: str) -> str:
    """Очищает ФИО преподавателя от служебных префиксов"""
    if not teacher_name:
        return None
    
    parts = teacher_name.split()
    if not parts:
        return None
    
    prefixes = ['преп.', 'ст.', 'доц.', 'проф.', 'асс.', ' зав.']
    cleaned_parts = [p for p in parts if p not in prefixes]
    
    return ' '.join(cleaned_parts) if cleaned_parts else None


def create_default_time_slots():
    """Создает типичные временные слоты"""
    default_slots = [
        ("1 пара", "09:00", "10:30"),
        ("2 пара", "10:40", "12:10"),
        ("3 пара", "12:20", "13:50"),
        ("4 пара", "14:00", "15:30"),
        ("5 пара", "15:40", "17:10"),
        ("6 пара", "17:20", "18:50"),
    ]
    
    for name, start, end in default_slots:
        from datetime import time
        s = time(*map(int, start.split(':')))
        e = time(*map(int, end.split(':')))
        TimeSlot.objects.get_or_create(
            name=name,
            defaults={'start_time': s, 'end_time': e}
        )


def _parse_lesson_type(text: str) -> str:
    """Определяет тип занятия"""
    if not text:
        return 'LEC'
    
    text_clean = re.sub(r'\s*\(.*?\)', '', text)
    text_lower = text_clean.lower()
    
    if 'прак' in text_lower or 'семинар' in text_lower:
        return 'PRA'
    if 'лаб' in text_lower:
        return 'LAB'
    return 'LEC'


def parse_all_tables_verbose(file_path: str):
    """Отладочная функция - показывает структуру всех таблиц"""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f'Файл не найден: {file_path}')
    
    doc = Document(file_path)
    result = {'tables': []}
    
    for table_idx, table in enumerate(doc.tables):
        table_info = {
            'idx': table_idx,
            'rows': len(table.rows),
            'cols': len(table.columns),
            'preview': []
        }
        
        for row_idx, row in enumerate(table.rows[:10]):
            cells_text = [cell.text.strip()[:50] for cell in row.cells]
            table_info['preview'].append(cells_text)
        
        result['tables'].append(table_info)
    
    return result